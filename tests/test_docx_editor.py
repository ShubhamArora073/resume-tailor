import os
import tempfile
import pytest
from docx import Document

from docx_editor import extract_paragraphs, apply_rewrites

FIXTURES_DIR = os.path.join(os.path.dirname(__file__), "fixtures")
SAMPLE_DOCX = os.path.join(FIXTURES_DIR, "sample.docx")


def create_test_docx(path: str, paragraphs: list[str]):
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def test_extract_paragraphs_returns_list():
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "test.docx")
        create_test_docx(docx_path, ["Hello World", "Second paragraph"])

        result = extract_paragraphs(docx_path)

        assert len(result) == 2
        assert result[0]["index"] == 0
        assert result[0]["text"] == "Hello World"
        assert result[1]["index"] == 1
        assert result[1]["text"] == "Second paragraph"


def test_extract_paragraphs_skips_empty():
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "test.docx")
        create_test_docx(docx_path, ["Content", "", "More content"])

        result = extract_paragraphs(docx_path)

        texts = [p["text"] for p in result]
        assert "" not in texts


def test_apply_rewrites_preserves_formatting():
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "test.docx")
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Original bold text")
        run.bold = True
        run.font.size = 140000  # 14pt
        doc.save(docx_path)

        output_path = os.path.join(tmp, "output.docx")
        rewrites = [{"index": 0, "rewritten": "Rewritten bold text"}]
        apply_rewrites(docx_path, rewrites, output_path)

        result_doc = Document(output_path)
        result_para = result_doc.paragraphs[0]
        assert result_para.text == "Rewritten bold text"
        assert result_para.runs[0].bold is True
        assert result_para.runs[0].font.size == 140000


def test_apply_rewrites_skips_missing_indices():
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "test.docx")
        create_test_docx(docx_path, ["Keep this", "Also keep"])

        output_path = os.path.join(tmp, "output.docx")
        rewrites = [{"index": 5, "rewritten": "Ghost"}]
        apply_rewrites(docx_path, rewrites, output_path)

        result_doc = Document(output_path)
        assert result_doc.paragraphs[0].text == "Keep this"
        assert result_doc.paragraphs[1].text == "Also keep"
