from docx import Document


def extract_paragraphs(docx_path: str) -> list[dict]:
    doc = Document(docx_path)
    paragraphs = []
    index = 0

    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        heading_level = None
        if para.style and para.style.name.startswith("Heading"):
            try:
                heading_level = int(para.style.name.split()[-1])
            except ValueError:
                heading_level = None

        paragraphs.append({
            "index": index,
            "text": para.text,
            "heading_level": heading_level,
        })
        index += 1

    return paragraphs


def apply_rewrites(docx_path: str, rewrites: list[dict], output_path: str) -> str:
    doc = Document(docx_path)

    rewrite_map = {r["index"]: r["rewritten"] for r in rewrites}

    non_empty_index = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue

        if non_empty_index in rewrite_map:
            new_text = rewrite_map[non_empty_index]
            if para.runs:
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = new_text

        non_empty_index += 1

    doc.save(output_path)
    return output_path
